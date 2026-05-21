"""Parse an uploaded document into raw text + paragraph blocks (spec §7, step 1).

I/O-pure: every parser takes the file's bytes and its suffix and returns a
`ParsedDocument`. No filesystem or network — the API layer reads the upload and
hands the bytes here. A `ParseError` (a `ValueError`) signals input the caller
should reject with a 400; it never leaks a library-specific exception type.

`.txt` and `.md` are treated identically: decode UTF-8, normalise line endings,
and split into blocks on blank lines (single newlines inside a block are kept,
so a wrapped paragraph stays one block). `.docx` is read with python-docx, one
block per non-empty paragraph.
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from io import BytesIO
from zipfile import BadZipFile

import docx
from docx.opc.exceptions import PackageNotFoundError

# Suffixes we accept, mapped to nothing — the set is the allowlist; the API layer
# validates the MIME/extension before calling, this is the second line of defence.
_TEXT_SUFFIXES = {".txt", ".md"}
_DOCX_SUFFIX = ".docx"

# A blank line is a run of two-or-more newlines (after CRLF normalisation),
# possibly with intervening whitespace; that is the paragraph-block separator.
_BLANK_LINE = re.compile(r"\n[ \t]*\n+")


class ParseError(ValueError):
    """Raised for input the caller should reject (unknown format, corrupt file)."""


@dataclass(frozen=True)
class ParsedDocument:
    """The result of parsing an upload: full text plus its paragraph blocks."""

    raw_text: str
    paragraphs: list[str]


def parse_document(data: bytes, suffix: str) -> ParsedDocument:
    """Parse `data` according to its file `suffix` (e.g. ".txt", ".docx")."""
    suffix = suffix.lower()
    if suffix in _TEXT_SUFFIXES:
        return _parse_text(data)
    if suffix == _DOCX_SUFFIX:
        return _parse_docx(data)
    raise ParseError(f"unsupported file type: {suffix!r}")


def _parse_text(data: bytes) -> ParsedDocument:
    try:
        text = data.decode("utf-8")
    except UnicodeDecodeError as exc:
        raise ParseError("file is not valid UTF-8 text") from exc
    raw_text = text.replace("\r\n", "\n").replace("\r", "\n")
    paragraphs = [block.strip() for block in _BLANK_LINE.split(raw_text)]
    return ParsedDocument(raw_text=raw_text, paragraphs=[p for p in paragraphs if p])


def _parse_docx(data: bytes) -> ParsedDocument:
    try:
        document = docx.Document(BytesIO(data))
    except (PackageNotFoundError, BadZipFile, ValueError) as exc:
        raise ParseError("file is not a readable .docx document") from exc
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    return ParsedDocument(raw_text="\n\n".join(paragraphs), paragraphs=paragraphs)
