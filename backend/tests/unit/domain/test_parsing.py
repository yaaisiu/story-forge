"""Unit tests for the document parsers (`domain/parsing.py`).

Pure functions: bytes + a file suffix in, a `ParsedDocument` (raw text + paragraph
blocks) out. No filesystem, no network — the docx fixture is built in memory.
"""

from __future__ import annotations

from io import BytesIO

import pytest
from docx import Document

from story_forge.domain.parsing import ParseError, parse_document


def test_txt_splits_on_blank_lines() -> None:
    text = "Para one, line one.\nPara one, line two.\n\nPara two.\n\n\nPara three.\n"
    result = parse_document(text.encode("utf-8"), ".txt")
    # A block is text between blank lines; single newlines inside a block are kept.
    assert result.paragraphs == [
        "Para one, line one.\nPara one, line two.",
        "Para two.",
        "Para three.",
    ]
    assert result.raw_text.startswith("Para one, line one.")
    assert "Para three." in result.raw_text


def test_md_parses_like_plain_text() -> None:
    md = "# Heading\n\nFirst body paragraph.\n\nSecond body paragraph.\n"
    result = parse_document(md.encode("utf-8"), ".md")
    assert result.paragraphs == ["# Heading", "First body paragraph.", "Second body paragraph."]


def test_crlf_line_endings_are_normalized() -> None:
    result = parse_document(b"Line A.\r\n\r\nLine B.\r\n", ".txt")
    assert result.paragraphs == ["Line A.", "Line B."]
    assert "\r" not in result.raw_text


def test_docx_uses_non_empty_paragraphs() -> None:
    buf = BytesIO()
    doc = Document()
    doc.add_paragraph("First paragraph.")
    doc.add_paragraph("")  # empty paragraph -> dropped
    doc.add_paragraph("Second paragraph.")
    doc.save(buf)

    result = parse_document(buf.getvalue(), ".docx")
    assert result.paragraphs == ["First paragraph.", "Second paragraph."]
    assert result.raw_text == "First paragraph.\n\nSecond paragraph."


def test_unknown_suffix_raises() -> None:
    with pytest.raises(ParseError):
        parse_document(b"whatever", ".pdf")


def test_invalid_utf8_text_raises() -> None:
    with pytest.raises(ParseError):
        parse_document(b"\xff\xfe not valid utf-8", ".txt")


def test_corrupt_docx_raises() -> None:
    with pytest.raises(ParseError):
        parse_document(b"not a real docx zip", ".docx")
